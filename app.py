# -*- coding: utf-8 -*-
"""
娜娜的工程质量监督Agent
=============================
功能：
1. 输入现场巡视问题描述，自动匹配相应的规章制度和条款
2. 内置规章制度数据库，支持导入更多规范
3. 生成标准格式的整改通知书

作者: WorkBuddy Agent
"""

import json
import os
import re
import sqlite3
import hashlib
import uuid
import threading
import time
from datetime import datetime
from flask import Flask, request, jsonify, render_template, send_file, Response

# LLM 适配器
from llm_adapter import LLMAdapter, PROVIDER_PRESETS

# 网络搜索与PDF解析
from web_search import search_legal_provisions, parse_search_results_to_regulation, fetch_page_text
from pdf_parser import extract_text_from_pdf, parse_pdf_text_to_regulation, ocr_pdf_images, parse_pdf_full_deep, smart_chunk_text

# ==============================================================================
# 路径配置
# ==============================================================================
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
DATA_DIR = os.path.join(BASE_DIR, "data")
DB_PATH = os.path.join(BASE_DIR, "regulations.db")
REGULATIONS_JSON = os.path.join(DATA_DIR, "regulations.json")

app = Flask(__name__, static_folder="static", template_folder="templates")
app.config["MAX_CONTENT_LENGTH"] = 50 * 1024 * 1024  # 限制上传50MB


# ==============================================================================
# PDF解析任务管理器（支持进度查询和终止）
# ==============================================================================
class PDFTaskManager:
    """
    管理PDF解析任务的异步执行、进度查询和终止。
    使用线程在后台执行解析，主线程可随时查询进度或请求终止。
    """
    def __init__(self):
        self._tasks = {}  # task_id -> task_info
        self._lock = threading.Lock()

    def create_task(self, file_bytes, filename, mode, use_ocr):
        """创建并启动一个PDF解析任务"""
        task_id = str(uuid.uuid4())[:8]
        cancel_event = threading.Event()

        task_info = {
            "task_id": task_id,
            "filename": filename,
            "mode": mode,
            "use_ocr": use_ocr,
            "status": "pending",  # pending / running / completed / failed / cancelled
            "progress": 0,
            "stage": "初始化中...",
            "result": None,
            "error": None,
            "cancel_event": cancel_event,
            "created_at": time.time(),
            "completed_at": None,
        }

        with self._lock:
            self._tasks[task_id] = task_info

        # 启动后台线程
        thread = threading.Thread(
            target=self._run_task,
            args=(task_id, file_bytes, mode, use_ocr, cancel_event),
            daemon=True,
        )
        thread.start()

        return task_id

    def _run_task(self, task_id, file_bytes, mode, use_ocr, cancel_event):
        """后台线程执行PDF解析"""
        task = self._tasks[task_id]

        def update_progress(progress, stage):
            with self._lock:
                task["progress"] = progress
                task["stage"] = stage
                task["status"] = "running"

        def is_cancelled():
            return cancel_event.is_set()

        def set_terminal_status(status, stage, result=None, error=None):
            """安全地设置终态，避免被worker覆盖"""
            with self._lock:
                if task["status"] == "cancelled":
                    return  # 已被取消，不覆盖
                task["status"] = status
                task["stage"] = stage
                if result is not None:
                    task["result"] = result
                if error is not None:
                    task["error"] = error
                task["completed_at"] = time.time()

        try:
            with self._lock:
                task["status"] = "running"
                task["stage"] = "正在读取PDF文件..."
                task["progress"] = 5

            if is_cancelled():
                set_terminal_status("cancelled", "用户已终止")
                return

            # ===== OCR 模式 =====
            if use_ocr:
                update_progress(10, "正在进行OCR识别（可能需要几分钟）...")
                ocr_result = ocr_pdf_images(file_bytes, max_pages=100)
                if is_cancelled():
                    set_terminal_status("cancelled", "用户已终止")
                    return

                if ocr_result.get("error"):
                    set_terminal_status("failed", "OCR识别失败", error=ocr_result["error"])
                    return

                text = ocr_result["text"]
                if not text or len(text.strip()) < 20:
                    set_terminal_status("failed", "OCR识别结果为空", error="OCR识别结果为空")
                    return

                update_progress(60, f"OCR完成，共{ocr_result['pages']}页，提取{len(text)}字符")

                if is_cancelled():
                    set_terminal_status("cancelled", "用户已终止")
                    return

                # OCR后深度解析
                if mode == "deep" and llm.is_ready():
                    update_progress(65, "AI深度解析中...")
                    parsed = parse_pdf_full_deep(
                        text, llm,
                        on_progress=lambda chunk_idx, total, clauses:
                            update_progress(
                                65 + int(30 * chunk_idx / max(total, 1)),
                                f"AI解析第{chunk_idx}/{total}块，已提取{clauses}条条款"
                            ),
                        cancel_check=is_cancelled,
                    )
                else:
                    parsed = parse_pdf_text_to_regulation(text, llm_adapter=llm)

                if is_cancelled():
                    set_terminal_status("cancelled", "用户已终止")
                    return

                reg = parsed.get("regulation", {})
                result = {
                    "message": f"OCR识别完成，共 {ocr_result['pages']} 页，提取文字 {len(text)} 字符"
                               + (f"，AI解析出 {parsed.get('clauses_count', len(reg.get('clauses', [])))} 条条款" if mode == "deep" else ""),
                    "text": text[:3000],
                    "pages": ocr_result["pages"],
                    "regulation": reg,
                    "raw_text": parsed.get("raw_text", ""),
                    "engine": "ocr" + ("+ai_deep" if mode == "deep" else ""),
                    "chunks": parsed.get("chunks", 1),
                    "clauses_count": parsed.get("clauses_count", len(reg.get("clauses", []))),
                }
                with self._lock:
                    task["progress"] = 100
                set_terminal_status("completed", "完成", result=result)
                return

            # ===== 正常文本提取 =====
            update_progress(15, "正在提取PDF文本...")
            extract_result = extract_text_from_pdf(file_bytes)

            if is_cancelled():
                set_terminal_status("cancelled", "用户已终止")
                return

            if extract_result.get("need_ocr"):
                result = {
                    "need_ocr": True,
                    "pages": extract_result.get("pages", 0),
                    "images": extract_result.get("images", []),
                }
                set_terminal_status("failed", "需要OCR识别", result=result,
                                    error=extract_result.get("error", "该PDF为扫描版，需要OCR识别"))
                return

            if extract_result.get("error") and not extract_result.get("text"):
                set_terminal_status("failed", "PDF文本提取失败", error=extract_result["error"])
                return

            text = extract_result["text"]
            pages = extract_result["pages"]
            engine = extract_result.get("engine", "")

            update_progress(35, f"文本提取完成（{pages}页，{len(text)}字符）")

            if is_cancelled():
                set_terminal_status("cancelled", "用户已终止")
                return

            # ===== 深度解析模式 =====
            if mode == "deep":
                if not llm.is_ready():
                    update_progress(50, "使用规则提取（未配置LLM）...")
                    parsed = parse_pdf_text_to_regulation(text, llm_adapter=None)
                    reg = parsed.get("regulation", {})
                    result = {
                        "message": f"PDF文本提取完成（{pages}页，{len(text)}字符），未配置LLM，使用规则提取到 {len(reg.get('clauses', []))} 条条款",
                        "text": text[:3000],
                        "pages": pages,
                        "regulation": reg,
                        "raw_text": parsed.get("raw_text", ""),
                        "engine": engine + "+rule",
                        "chunks": 1,
                        "clauses_count": len(reg.get("clauses", [])),
                        "warning": "未配置大模型API，仅使用规则提取。",
                    }
                    with self._lock:
                        task["progress"] = 100
                    set_terminal_status("completed", "完成（规则提取）", result=result)
                    return

                # AI深度解析
                update_progress(40, "AI深度解析中...")
                chunks_info = smart_chunk_text(text)
                total_chunks = len(chunks_info)

                parsed = parse_pdf_full_deep(
                    text, llm,
                    on_progress=lambda chunk_idx, total, clauses:
                        update_progress(
                            40 + int(55 * chunk_idx / max(total, 1)),
                            f"AI解析第{chunk_idx}/{total}块，已提取{clauses}条条款"
                        ),
                    cancel_check=is_cancelled,
                )

                if is_cancelled():
                    set_terminal_status("cancelled", "用户已终止")
                    return

                reg = parsed.get("regulation", {})
                clauses_count = parsed.get("clauses_count", len(reg.get("clauses", [])))
                errors = parsed.get("errors", [])

                msg = (f"PDF深度解析完成：{pages}页，{len(text)}字符，分{total_chunks}块处理，"
                       f"AI提取出 {clauses_count} 条条款")

                result = {
                    "message": msg,
                    "text": text[:3000],
                    "pages": pages,
                    "regulation": reg,
                    "raw_text": parsed.get("raw_text", ""),
                    "engine": engine + "+ai_deep",
                    "chunks": total_chunks,
                    "clauses_count": clauses_count,
                    "full_text_length": len(text),
                }
                if errors:
                    result["warnings"] = errors[:5]

                with self._lock:
                    task["progress"] = 100
                set_terminal_status("completed", "完成", result=result)
                return

            # ===== 普通模式 =====
            update_progress(60, "解析规范信息中...")
            parsed = parse_pdf_text_to_regulation(text, llm_adapter=llm)
            reg = parsed.get("regulation", {})

            if is_cancelled():
                set_terminal_status("cancelled", "用户已终止")
                return

            result = {
                "message": f"PDF解析完成，共 {pages} 页（引擎: {engine}）",
                "text": text[:3000],
                "pages": pages,
                "regulation": reg,
                "raw_text": parsed.get("raw_text", ""),
                "engine": engine,
                "clauses_count": len(reg.get("clauses", [])),
                "full_text_length": len(text),
            }
            with self._lock:
                task["progress"] = 100
            set_terminal_status("completed", "完成", result=result)

        except Exception as e:
            set_terminal_status("failed", "解析异常", error=f"解析异常: {str(e)}")

    def get_task(self, task_id):
        """获取任务状态"""
        with self._lock:
            task = self._tasks.get(task_id)
            if not task:
                return None
            return {
                "task_id": task["task_id"],
                "filename": task["filename"],
                "mode": task["mode"],
                "status": task["status"],
                "progress": task["progress"],
                "stage": task["stage"],
                "result": task["result"],
                "error": task["error"],
                "elapsed": round(time.time() - task["created_at"], 1),
            }

    def cancel_task(self, task_id):
        """请求终止任务"""
        with self._lock:
            task = self._tasks.get(task_id)
            if not task:
                return False
            if task["status"] in ("completed", "failed", "cancelled"):
                return False
            task["cancel_event"].set()
            task["status"] = "cancelled"
            task["stage"] = "用户已终止"
            return True

    def cleanup_old_tasks(self, max_age=3600):
        """清理超过1小时的旧任务"""
        with self._lock:
            now = time.time()
            to_remove = [
                tid for tid, t in self._tasks.items()
                if t.get("completed_at") and now - t["completed_at"] > max_age
            ]
            for tid in to_remove:
                del self._tasks[tid]


pdf_task_manager = PDFTaskManager()


# ==============================================================================
# 数据库管理
# ==============================================================================
class RegulationDB:
    """规章制度数据库管理类"""

    def __init__(self, db_path):
        self.db_path = db_path
        self._init_db()

    def _get_conn(self):
        conn = sqlite3.connect(self.db_path)
        conn.row_factory = sqlite3.Row
        conn.execute("PRAGMA foreign_keys = ON")
        return conn

    def _init_db(self):
        conn = self._get_conn()
        conn.executescript("""
            CREATE TABLE IF NOT EXISTS regulations (
                id          INTEGER PRIMARY KEY AUTOINCREMENT,
                name        TEXT NOT NULL,
                code        TEXT NOT NULL UNIQUE,
                full_name   TEXT,
                category    TEXT,
                publish_date TEXT,
                implement_date TEXT,
                is_mandatory INTEGER DEFAULT 0,
                description TEXT,
                created_at  TEXT DEFAULT (datetime('now','localtime'))
            );

            CREATE TABLE IF NOT EXISTS clauses (
                id            INTEGER PRIMARY KEY AUTOINCREMENT,
                regulation_id INTEGER NOT NULL,
                clause_number TEXT NOT NULL,
                clause_content TEXT NOT NULL,
                keywords      TEXT,
                category      TEXT,
                is_mandatory  INTEGER DEFAULT 0,
                FOREIGN KEY (regulation_id) REFERENCES regulations(id) ON DELETE CASCADE
            );

            CREATE TABLE IF NOT EXISTS problems (
                id            INTEGER PRIMARY KEY AUTOINCREMENT,
                description   TEXT NOT NULL,
                location      TEXT,
                project_name  TEXT,
                inspection_date TEXT,
                matched_clauses TEXT,
                created_at    TEXT DEFAULT (datetime('now','localtime'))
            );

            CREATE TABLE IF NOT EXISTS notice_records (
                id            INTEGER PRIMARY KEY AUTOINCREMENT,
                project_name  TEXT,
                supervision_no TEXT,
                construction_unit TEXT,
                supervision_unit TEXT,
                construction_company TEXT,
                inspection_date TEXT,
                problems      TEXT,
                created_at    TEXT DEFAULT (datetime('now','localtime'))
            );
        """)
        conn.commit()
        conn.close()

    def load_from_json(self, json_path):
        """从JSON文件加载规章制度数据"""
        with open(json_path, "r", encoding="utf-8") as f:
            data = json.load(f)

        conn = self._get_conn()
        count = 0
        for reg in data.get("regulations", []):
            # 检查是否已存在
            existing = conn.execute(
                "SELECT id FROM regulations WHERE code = ?", (reg["code"],)
            ).fetchone()

            if existing:
                reg_id = existing["id"]
            else:
                cur = conn.execute(
                    """INSERT INTO regulations
                       (name, code, full_name, category, publish_date,
                        implement_date, is_mandatory, description)
                       VALUES (?, ?, ?, ?, ?, ?, ?, ?)""",
                    (
                        reg["name"], reg["code"], reg.get("full_name", ""),
                        reg.get("category", ""), reg.get("publish_date", ""),
                        reg.get("implement_date", ""),
                        1 if reg.get("is_mandatory") else 0,
                        reg.get("description", ""),
                    ),
                )
                reg_id = cur.lastrowid

            # 插入条款
            for clause in reg.get("clauses", []):
                # 检查条款是否已存在
                existing_clause = conn.execute(
                    "SELECT id FROM clauses WHERE regulation_id = ? AND clause_number = ?",
                    (reg_id, clause["clause_number"]),
                ).fetchone()

                if not existing_clause:
                    keywords_str = ",".join(clause.get("keywords", []))
                    conn.execute(
                        """INSERT INTO clauses
                           (regulation_id, clause_number, clause_content,
                            keywords, category, is_mandatory)
                           VALUES (?, ?, ?, ?, ?, ?)""",
                        (
                            reg_id, clause["clause_number"],
                            clause["clause_content"], keywords_str,
                            clause.get("category", ""), 0,
                        ),
                    )
                    count += 1

        conn.commit()
        conn.close()
        return count

    def get_all_regulations(self):
        """获取所有规章制度"""
        conn = self._get_conn()
        rows = conn.execute(
            "SELECT * FROM regulations ORDER BY category, code"
        ).fetchall()
        conn.close()
        return [dict(r) for r in rows]

    def get_regulation_detail(self, reg_id):
        """获取单个规章制度的详细信息（含条款）"""
        conn = self._get_conn()
        reg = conn.execute(
            "SELECT * FROM regulations WHERE id = ?", (reg_id,)
        ).fetchone()
        if not reg:
            conn.close()
            return None
        clauses = conn.execute(
            "SELECT * FROM clauses WHERE regulation_id = ? ORDER BY clause_number",
            (reg_id,),
        ).fetchall()
        conn.close()
        return {"regulation": dict(reg), "clauses": [dict(c) for c in clauses]}

    def add_regulation(self, data):
        """添加新规章制度"""
        conn = self._get_conn()
        try:
            cur = conn.execute(
                """INSERT INTO regulations
                   (name, code, full_name, category, publish_date,
                    implement_date, is_mandatory, description)
                   VALUES (?, ?, ?, ?, ?, ?, ?, ?)""",
                (
                    data["name"], data["code"], data.get("full_name", ""),
                    data.get("category", ""), data.get("publish_date", ""),
                    data.get("implement_date", ""),
                    1 if data.get("is_mandatory") else 0,
                    data.get("description", ""),
                ),
            )
            reg_id = cur.lastrowid

            for clause in data.get("clauses", []):
                keywords_str = ",".join(clause.get("keywords", []))
                conn.execute(
                    """INSERT INTO clauses
                       (regulation_id, clause_number, clause_content,
                        keywords, category, is_mandatory)
                       VALUES (?, ?, ?, ?, ?, ?)""",
                    (
                        reg_id, clause["clause_number"], clause["clause_content"],
                        keywords_str, clause.get("category", ""),
                        1 if clause.get("is_mandatory") else 0,
                    ),
                )
            conn.commit()
            conn.close()
            return reg_id
        except sqlite3.IntegrityError:
            conn.close()
            return None

    def import_regulations(self, json_data):
        """批量导入规章制度"""
        if isinstance(json_data, str):
            data = json.loads(json_data)
        else:
            data = json_data

        regs = data.get("regulations", data) if isinstance(data, dict) else data
        conn = self._get_conn()
        count = 0
        for reg in regs:
            existing = conn.execute(
                "SELECT id FROM regulations WHERE code = ?", (reg["code"],)
            ).fetchone()
            if existing:
                continue
            cur = conn.execute(
                """INSERT INTO regulations
                   (name, code, full_name, category, publish_date,
                    implement_date, is_mandatory, description)
                   VALUES (?, ?, ?, ?, ?, ?, ?, ?)""",
                (
                    reg["name"], reg["code"], reg.get("full_name", ""),
                    reg.get("category", ""), reg.get("publish_date", ""),
                    reg.get("implement_date", ""),
                    1 if reg.get("is_mandatory") else 0,
                    reg.get("description", ""),
                ),
            )
            reg_id = cur.lastrowid
            for clause in reg.get("clauses", []):
                keywords_str = ",".join(clause.get("keywords", []))
                conn.execute(
                    """INSERT INTO clauses
                       (regulation_id, clause_number, clause_content,
                        keywords, category, is_mandatory)
                       VALUES (?, ?, ?, ?, ?, ?)""",
                    (
                        reg_id, clause["clause_number"], clause["clause_content"],
                        keywords_str, clause.get("category", ""),
                        1 if clause.get("is_mandatory") else 0,
                    ),
                )
            count += 1
        conn.commit()
        conn.close()
        return count

    def get_all_clauses(self):
        """获取所有条款（用于匹配）"""
        conn = self._get_conn()
        rows = conn.execute(
            """SELECT c.*, r.name as reg_name, r.code as reg_code,
                      r.full_name as reg_full_name, r.is_mandatory as reg_mandatory
               FROM clauses c
               JOIN regulations r ON c.regulation_id = r.id
               ORDER BY r.category, r.code"""
        ).fetchall()
        conn.close()
        return [dict(r) for r in rows]

    def delete_regulation(self, reg_id):
        """删除整部规章制度（含其下所有条款）"""
        conn = self._get_conn()
        cur = conn.execute("DELETE FROM regulations WHERE id = ?", (reg_id,))
        deleted = cur.rowcount
        # 外键级联删除会自动删除 clauses
        conn.commit()
        conn.close()
        return deleted

    def delete_clause(self, clause_id):
        """删除单条条款"""
        conn = self._get_conn()
        cur = conn.execute("DELETE FROM clauses WHERE id = ?", (clause_id,))
        deleted = cur.rowcount
        conn.commit()
        conn.close()
        return deleted

    def batch_delete_regulations(self, reg_ids):
        """批量删除多部规章制度"""
        conn = self._get_conn()
        count = 0
        for reg_id in reg_ids:
            cur = conn.execute("DELETE FROM regulations WHERE id = ?", (reg_id,))
            count += cur.rowcount
        conn.commit()
        conn.close()
        return count

    def batch_delete_clauses(self, clause_ids):
        """批量删除多条条款"""
        conn = self._get_conn()
        count = 0
        for clause_id in clause_ids:
            cur = conn.execute("DELETE FROM clauses WHERE id = ?", (clause_id,))
            count += cur.rowcount
        conn.commit()
        conn.close()
        return count

    def batch_add_regulations(self, regulations_list):
        """批量添加多部规章制度（跳过已存在的编号）"""
        conn = self._get_conn()
        added = 0
        skipped = 0
        for reg in regulations_list:
            code = reg.get("code", "")
            if not code:
                skipped += 1
                continue
            existing = conn.execute(
                "SELECT id FROM regulations WHERE code = ?", (code,)
            ).fetchone()
            if existing:
                skipped += 1
                continue
            try:
                cur = conn.execute(
                    """INSERT INTO regulations
                       (name, code, full_name, category, publish_date,
                        implement_date, is_mandatory, description)
                       VALUES (?, ?, ?, ?, ?, ?, ?, ?)""",
                    (
                        reg.get("name", ""), code, reg.get("full_name", ""),
                        reg.get("category", ""), reg.get("publish_date", ""),
                        reg.get("implement_date", ""),
                        1 if reg.get("is_mandatory") else 0,
                        reg.get("description", ""),
                    ),
                )
                reg_id = cur.lastrowid
                for clause in reg.get("clauses", []):
                    keywords_str = ",".join(clause.get("keywords", []))
                    conn.execute(
                        """INSERT INTO clauses
                           (regulation_id, clause_number, clause_content,
                            keywords, category, is_mandatory)
                           VALUES (?, ?, ?, ?, ?, ?)""",
                        (
                            reg_id, clause.get("clause_number", ""),
                            clause.get("clause_content", ""),
                            keywords_str, clause.get("category", ""),
                            1 if clause.get("is_mandatory") else 0,
                        ),
                    )
                added += 1
            except Exception:
                skipped += 1
        conn.commit()
        conn.close()
        return {"added": added, "skipped": skipped}

    def update_regulation(self, reg_id, data):
        """更新规章制度信息"""
        conn = self._get_conn()
        fields = []
        values = []
        for key in ("name", "code", "full_name", "category", "publish_date",
                     "implement_date", "is_mandatory", "description"):
            if key in data:
                fields.append(f"{key} = ?")
                val = data[key]
                if key == "is_mandatory":
                    val = 1 if val else 0
                values.append(val)
        if not fields:
            conn.close()
            return 0
        values.append(reg_id)
        cur = conn.execute(
            f"UPDATE regulations SET {', '.join(fields)} WHERE id = ?", values
        )
        conn.commit()
        deleted = cur.rowcount
        conn.close()
        return deleted

    def add_clause(self, reg_id, clause_data):
        """向已有规范添加单条条款"""
        conn = self._get_conn()
        keywords_str = ",".join(clause_data.get("keywords", []))
        cur = conn.execute(
            """INSERT INTO clauses
               (regulation_id, clause_number, clause_content,
                keywords, category, is_mandatory)
               VALUES (?, ?, ?, ?, ?, ?)""",
            (
                reg_id, clause_data["clause_number"], clause_data["clause_content"],
                keywords_str, clause_data.get("category", ""),
                1 if clause_data.get("is_mandatory") else 0,
            ),
        )
        clause_id = cur.lastrowid
        conn.commit()
        conn.close()
        return clause_id

    def get_stats(self):
        """获取数据库统计信息"""
        conn = self._get_conn()
        reg_count = conn.execute("SELECT COUNT(*) as c FROM regulations").fetchone()["c"]
        clause_count = conn.execute("SELECT COUNT(*) as c FROM clauses").fetchone()["c"]
        categories = conn.execute(
            "SELECT DISTINCT category FROM regulations WHERE category != ''"
        ).fetchall()
        conn.close()
        return {
            "regulations": reg_count,
            "clauses": clause_count,
            "categories": [c["category"] for c in categories],
        }


# ==============================================================================
# 智能匹配引擎
# ==============================================================================
class RegulationMatcher:
    """问题-规范智能匹配引擎"""

    # 领域同义词映射（用于扩展匹配）
    SYNONYMS = {
        "消火栓": ["消防栓", "消火栓"],
        "防火封堵": ["封堵", "防火封堵", "封堵材料"],
        "应急照明": ["应急灯", "应急照明", "疏散照明"],
        "防火保护": ["防火措施", "防火保护", "防火涂层"],
        "防火窗": ["防火玻璃窗", "防火窗"],
        "挡烟垂壁": ["挡烟垂壁", "挡烟壁", "垂壁"],
        "减压栓": ["减压消火栓", "减压栓", "减压设施"],
        "水龙带": ["水带", "水龙带", "消防水带"],
        "阻燃管": ["阻燃管", "阻燃套管", "穿管保护"],
        "疏散门": ["疏散门", "安全门", "逃生门"],
        "管道支架": ["支架", "管架", "管道支架", "支吊架"],
        "防火分隔": ["防火分隔", "防火隔离", "分隔"],
        "疏散楼梯间": ["楼梯间", "疏散楼梯", "安全楼梯"],
        "防火隔墙": ["防火墙", "防火隔墙", "隔火墙"],
        "桥架": ["电缆桥架", "桥架"],
    }

    # 问题类型关键词映射到规范类别
    CATEGORY_KEYWORDS = {
        "消防防火": ["防火", "消防", "耐火", "燃烧", "可燃", "不燃", "阻燃"],
        "消防给水": ["消火栓", "消防水", "消防给水", "水龙带", "消防泵", "消防水箱", "减压栓"],
        "消防电气": ["应急照明", "消防电缆", "消防配电", "应急灯", "火报", "火灾报警"],
        "防烟排烟": ["挡烟", "排烟", "防烟", "送风", "烟气"],
        "疏散设施": ["疏散", "楼梯间", "安全出口", "疏散门", "疏散通道"],
        "建筑电气": ["线缆", "电缆", "电线", "穿管", "过线盒", "电气线路", "配电"],
        "结构安全": ["混凝土", "钢筋", "砌体", "裂缝", "蜂窝", "露筋", "钢结构", "螺栓", "焊接"],
        "给排水": ["给水", "排水", "管道", "水压", "坡度"],
        "建筑节能": ["保温", "节能", "隔热", "气密"],
    }

    def __init__(self, db):
        self.db = db
        self._clauses_cache = None

    def _get_clauses(self):
        if self._clauses_cache is None:
            self._clauses_cache = self.db.get_all_clauses()
        return self._clauses_cache

    def _extract_ngrams(self, text, n_list=(2, 3, 4)):
        """提取文本中的n-gram序列"""
        text = re.sub(r'[^\u4e00-\u9fa5a-zA-Z0-9]', '', text)
        ngrams = set()
        for n in n_list:
            for i in range(len(text) - n + 1):
                ngrams.add(text[i:i + n])
        return ngrams

    def _expand_keyword(self, keyword):
        """扩展关键词（包含同义词）"""
        expanded = {keyword}
        for key, synonyms in self.SYNONYMS.items():
            if keyword in synonyms or keyword == key:
                expanded.update(synonyms)
        return expanded

    def _detect_categories(self, text):
        """检测问题文本所属的类别"""
        detected = set()
        for category, keywords in self.CATEGORY_KEYWORDS.items():
            for kw in keywords:
                if kw in text:
                    detected.add(category)
                    break
        return detected

    def match(self, problem_text, top_n=5):
        """
        匹配核心方法：输入问题描述，返回匹配的规范条款

        参数:
            problem_text: 问题描述文本
            top_n: 返回的最佳匹配数量

        返回:
            排序后的匹配结果列表
        """
        if not problem_text or not problem_text.strip():
            return []

        clauses = self._get_clauses()
        if not clauses:
            return []

        # 检测问题类别
        detected_categories = self._detect_categories(problem_text)

        # 提取问题的n-gram
        problem_ngrams = self._extract_ngrams(problem_text)

        results = []

        for clause in clauses:
            score = 0
            matched_keywords = []
            match_details = []

            # 获取条款关键词
            keywords = []
            if clause.get("keywords"):
                keywords = [k.strip() for k in clause["keywords"].split(",") if k.strip()]

            # 关键词匹配
            for kw in keywords:
                expanded_kws = self._expand_keyword(kw)
                for ekw in expanded_kws:
                    if ekw in problem_text:
                        # 长关键词权重更高
                        kw_score = len(ekw) * 2
                        if len(ekw) >= 4:
                            kw_score += 3
                        score += kw_score
                        matched_keywords.append(kw)
                        match_details.append(f"关键词匹配: {ekw}")
                        break

            # n-gram匹配（条款内容）
            clause_content = clause.get("clause_content", "")
            clause_ngrams = self._extract_ngrams(clause_content)
            common_ngrams = problem_ngrams & clause_ngrams

            # 过滤掉太短的通用n-gram
            meaningful_common = [ng for ng in common_ngrams if len(ng) >= 3]
            if meaningful_common:
                ngram_score = len(meaningful_common) * 2
                score += ngram_score
                match_details.append(f"内容匹配: {len(meaningful_common)}个共同词组")

            # 类别匹配加分
            clause_category = clause.get("category", "")
            if clause_category and clause_category in detected_categories:
                score += 5
                match_details.append(f"类别匹配: {clause_category}")

            # 规范类别匹配
            reg_category = None
            for cat in self.CATEGORY_KEYWORDS:
                if cat in (clause.get("category", ""),):
                    reg_category = cat
                    break

            if score > 0:
                results.append({
                    "clause_id": clause["id"],
                    "regulation_name": clause["reg_name"],
                    "regulation_code": clause["reg_code"],
                    "regulation_full_name": clause.get("reg_full_name", clause["reg_name"]),
                    "is_mandatory_reg": bool(clause.get("reg_mandatory", 0)),
                    "clause_number": clause["clause_number"],
                    "clause_content": clause["clause_content"],
                    "clause_category": clause.get("category", ""),
                    "keywords": keywords,
                    "matched_keywords": list(set(matched_keywords)),
                    "match_score": score,
                    "match_details": match_details,
                })

        # 按匹配分数排序
        results.sort(key=lambda x: x["match_score"], reverse=True)

        # 归一化匹配分数到0-100
        if results:
            max_score = max(r["match_score"] for r in results)
            for r in results:
                r["match_percentage"] = round(min(100, r["match_score"] / max_score * 100)) if max_score > 0 else 0
                # 生成匹配等级
                if r["match_percentage"] >= 80:
                    r["match_level"] = "高"
                elif r["match_percentage"] >= 50:
                    r["match_level"] = "中"
                else:
                    r["match_level"] = "低"

        return results[:top_n]

    def generate_notice_text(self, problems, project_info):
        """
        生成整改通知书文本

        参数:
            problems: 问题列表，每个问题包含description和matched结果
            project_info: 工程信息字典
        """
        lines = []
        lines.append("工程质量监督抽查整改通知书")
        lines.append(f"  监督注册号：{project_info.get('supervision_no', '')}")
        lines.append("")
        lines.append(f"{project_info.get('construction_unit', '')}：")
        lines.append(f"{project_info.get('supervision_unit', '')}：")
        lines.append(f"{project_info.get('construction_company', '')}：")
        lines.append("")

        inspection_date = project_info.get('inspection_date', '')
        project_name = project_info.get('project_name', '')

        lines.append(f"经我站{inspection_date}监督抽查，发现 {project_name} 存在以下质量问题：")
        lines.append("")

        for i, prob in enumerate(problems, 1):
            desc = prob["description"]
            matches = prob.get("matches", [])

            if matches:
                # 取第一个最佳匹配
                best = matches[0]
                refs = []
                for m in matches[:3]:  # 最多引用3条
                    ref = f"《{m['regulation_full_name']}》{m['regulation_code']}第{m['clause_number']}条"
                    refs.append(ref)
                ref_text = "、".join(refs)
                lines.append(f"{i}. {desc}，不符合{ref_text}。")
            else:
                lines.append(f"{i}. {desc}。")
            lines.append("")

        lines.append("请你单位整改完毕后并经建设或监理单位检查合格并签署意见后，")
        lines.append("将整改完成报告书面报我站，整改完成证明文件及资料（影像资料）")
        lines.append("由建设、监理、施工单位存档备查。")
        lines.append("")
        lines.append("特此通知！")
        lines.append("")
        lines.append(f"质量监督员：　　　　　　　")
        lines.append(f"建设工程质量监督站")
        lines.append(f"  （盖章）")
        lines.append(f"  {inspection_date}")
        lines.append("")
        lines.append("签收人：    年  月  日")

        return "\n".join(lines)


# ==============================================================================
# 初始化
# ==============================================================================
db = RegulationDB(DB_PATH)
matcher = RegulationMatcher(db)
llm = LLMAdapter()

# 自动加载预置数据
if os.path.exists(REGULATIONS_JSON):
    loaded = db.load_from_json(REGULATIONS_JSON)
    if loaded > 0:
        print(f"[INFO] 预加载 {loaded} 条规范条款")
    else:
        print("[INFO] 规范数据已存在，无需重复加载")

if llm.is_ready():
    print(f"[INFO] LLM已启用: {llm.config['provider']} / {llm.config['model']}")
else:
    print("[INFO] LLM未启用（可在设置中配置）")


# ==============================================================================
# API 路由
# ==============================================================================
@app.route("/")
def index():
    return render_template("index.html")


@app.route("/api/match", methods=["POST"])
def api_match():
    """匹配API：输入问题描述，返回匹配的规范条款"""
    data = request.get_json()
    if not data or "problem" not in data:
        return jsonify({"error": "请提供问题描述"}), 400

    problem = data["problem"].strip()
    location = data.get("location", "")
    top_n = data.get("top_n", 5)

    if location:
        full_problem = f"{problem}（位置：{location}）"
    else:
        full_problem = problem

    matches = matcher.match(full_problem, top_n=top_n)

    return jsonify({
        "problem": problem,
        "location": location,
        "matches": matches,
        "total": len(matches),
    })


@app.route("/api/match/batch", methods=["POST"])
def api_match_batch():
    """批量匹配API：输入多个问题，返回每个问题的匹配结果"""
    data = request.get_json()
    if not data or "problems" not in data:
        return jsonify({"error": "请提供问题列表"}), 400

    problems = data["problems"]
    results = []

    for prob in problems:
        if isinstance(prob, str):
            desc = prob
            location = ""
        else:
            desc = prob.get("description", "")
            location = prob.get("location", "")

        full_problem = f"{desc}（位置：{location}）" if location else desc
        matches = matcher.match(full_problem, top_n=5)
        results.append({
            "description": desc,
            "location": location,
            "matches": matches,
        })

    return jsonify({"results": results, "total": len(results)})


@app.route("/api/regulations", methods=["GET"])
def api_get_regulations():
    """获取所有规章制度列表"""
    regs = db.get_all_regulations()
    stats = db.get_stats()
    return jsonify({"regulations": regs, "stats": stats})


@app.route("/api/regulations/<int:reg_id>", methods=["GET"])
def api_get_regulation_detail(reg_id):
    """获取单个规章制度详情"""
    detail = db.get_regulation_detail(reg_id)
    if not detail:
        return jsonify({"error": "未找到该规章制度"}), 404
    return jsonify(detail)


@app.route("/api/regulations/import", methods=["POST"])
def api_import_regulations():
    """导入规章制度"""
    data = request.get_json()
    if not data:
        return jsonify({"error": "请提供规章制度数据"}), 400

    count = db.import_regulations(data)
    matcher._clauses_cache = None  # 清除缓存
    return jsonify({"message": f"成功导入 {count} 部规章制度", "count": count})


@app.route("/api/regulations/add", methods=["POST"])
def api_add_regulation():
    """添加单个规章制度"""
    data = request.get_json()
    if not data or "name" not in data or "code" not in data:
        return jsonify({"error": "请提供规范名称和编号"}), 400

    reg_id = db.add_regulation(data)
    if reg_id is None:
        return jsonify({"error": "该规范编号已存在"}), 409

    matcher._clauses_cache = None
    return jsonify({"message": "添加成功", "id": reg_id}), 201


@app.route("/api/regulations/<int:reg_id>", methods=["DELETE"])
def api_delete_regulation(reg_id):
    """删除整部规章制度"""
    deleted = db.delete_regulation(reg_id)
    if deleted == 0:
        return jsonify({"error": "未找到该规章制度"}), 404
    matcher._clauses_cache = None
    return jsonify({"message": "已删除该规章制度及其所有条款"})


@app.route("/api/regulations/<int:reg_id>/clauses/<int:clause_id>", methods=["DELETE"])
def api_delete_clause(reg_id, clause_id):
    """删除单条条款"""
    deleted = db.delete_clause(clause_id)
    if deleted == 0:
        return jsonify({"error": "未找到该条款"}), 404
    matcher._clauses_cache = None
    return jsonify({"message": "已删除该条款"})


@app.route("/api/regulations/<int:reg_id>", methods=["PUT"])
def api_update_regulation(reg_id):
    """更新规章制度信息"""
    data = request.get_json()
    if not data:
        return jsonify({"error": "请提供更新数据"}), 400
    updated = db.update_regulation(reg_id, data)
    if updated == 0:
        return jsonify({"error": "未找到该规章制度或无更新字段"}), 404
    matcher._clauses_cache = None
    return jsonify({"message": "更新成功"})


@app.route("/api/regulations/<int:reg_id>/clauses", methods=["POST"])
def api_add_clause(reg_id):
    """向已有规范添加单条条款"""
    data = request.get_json()
    if not data or "clause_number" not in data or "clause_content" not in data:
        return jsonify({"error": "请提供条款号和条款内容"}), 400
    clause_id = db.add_clause(reg_id, data)
    matcher._clauses_cache = None
    return jsonify({"message": "条款添加成功", "id": clause_id}), 201


@app.route("/api/regulations/batch_delete", methods=["POST"])
def api_batch_delete_regulations():
    """批量删除规章制度"""
    data = request.get_json()
    if not data or "ids" not in data:
        return jsonify({"error": "请提供要删除的规范ID列表"}), 400

    ids = data["ids"]
    if not isinstance(ids, list) or len(ids) == 0:
        return jsonify({"error": "ID列表不能为空"}), 400

    deleted = db.batch_delete_regulations(ids)
    matcher._clauses_cache = None
    return jsonify({"message": f"已删除 {deleted} 部规章制度", "deleted": deleted})


@app.route("/api/regulations/batch_add", methods=["POST"])
def api_batch_add_regulations():
    """批量添加规章制度"""
    data = request.get_json()
    if not data:
        return jsonify({"error": "请提供规范数据"}), 400

    # 兼容 {regulations: [...]} 或 直接 [...]
    regs = data.get("regulations", data) if isinstance(data, dict) else data
    if not isinstance(regs, list):
        return jsonify({"error": "数据格式错误，应为数组"}), 400

    result = db.batch_add_regulations(regs)
    matcher._clauses_cache = None
    return jsonify({
        "message": f"成功添加 {result['added']} 部规范，跳过 {result['skipped']} 部（已存在或格式错误）",
        "added": result["added"],
        "skipped": result["skipped"],
    })


@app.route("/api/regulations/search", methods=["POST"])
def api_search_regulations():
    """在线搜索法律/规范条文"""
    data = request.get_json()
    if not data or "query" not in data:
        return jsonify({"error": "请提供搜索关键词"}), 400

    query = data["query"].strip()
    if not query:
        return jsonify({"error": "搜索关键词不能为空"}), 400

    max_results = data.get("max_results", 5)
    search_result = search_legal_provisions(query, max_results=max_results)

    if not search_result.get("results"):
        return jsonify({
            "message": "未获取到搜索结果，请尝试更换关键词或检查网络",
            "results": [],
            "source": search_result.get("source", ""),
            "regulation": None,
        })

    parsed = parse_search_results_to_regulation(
        search_result["results"], query, llm_adapter=llm
    )

    return jsonify({
        "message": f"从 {search_result.get('source', '网络')} 搜索到 {len(search_result['results'])} 条结果",
        "results": search_result["results"],
        "source": search_result.get("source", ""),
        "regulation": parsed.get("regulation"),
        "raw_text": parsed.get("raw_text", ""),
    })


@app.route("/api/regulations/fetch_url", methods=["POST"])
def api_fetch_url():
    """抓取指定URL页面文本"""
    data = request.get_json()
    if not data or "url" not in data:
        return jsonify({"error": "请提供URL"}), 400

    url = data["url"].strip()
    text = fetch_page_text(url, max_chars=20000)
    if not text:
        return jsonify({"error": "无法获取页面内容"}), 500

    # 如果配置了LLM，自动解析全文为结构化规范数据
    parsed = parse_search_results_to_regulation(
        [{"title": url, "snippet": text[:500], "url": url}],
        url, llm_adapter=llm
    )

    # 用全文而非摘要作为 raw_text
    full_raw = f"来源URL：{url}\n\n全文内容：\n{text}"

    # 如果LLM解析成功，用全文重新解析以获取更多条款
    if llm.is_ready() and text:
        regulation = _parse_full_text_with_llm(text, url, llm)
        if regulation:
            parsed = {"regulation": regulation, "raw_text": full_raw}

    return jsonify({
        "text": text[:3000],
        "full_text_length": len(text),
        "regulation": parsed.get("regulation", {}),
        "raw_text": full_raw,
    })


def _parse_full_text_with_llm(full_text, source_url, llm_adapter):
    """使用LLM解析全文为完整规范JSON（含所有条款）"""
    # 截取前12000字符（约6000汉字），覆盖大部分规范条文
    text_chunk = full_text[:12000]

    system_prompt = (
        "你是中国工程建设标准领域的专家。请从以下网页全文中提取完整的规范信息。"
        "尽量提取所有可识别的条款（条文），不要遗漏。"
        "如果文本包含完整的规范条文，请逐条提取。"
    )
    user_prompt = (
        f"来源：{source_url}\n\n{text_chunk}\n\n"
        "请按以下JSON格式返回（只返回JSON，不要其他内容）：\n"
        "{\n"
        '  "name": "规范名称",\n'
        '  "code": "规范编号",\n'
        '  "full_name": "规范全称",\n'
        '  "category": "分类",\n'
        '  "is_mandatory": true或false,\n'
        '  "description": "规范简介",\n'
        '  "clauses": [\n'
        '    {"clause_number": "条款号", "clause_content": "条款全文内容", "keywords": ["关键词"]}\n'
        '  ]\n'
        "}\n"
    )
    messages = [
        {"role": "system", "content": system_prompt},
        {"role": "user", "content": user_prompt},
    ]
    result = llm_adapter.chat(messages, temperature=0.2, max_tokens=8192)
    if "error" in result:
        return None

    content = result["content"].strip()
    if content.startswith("```"):
        lines = content.split("\n")
        lines = [l for l in lines if not l.strip().startswith("```")]
        content = "\n".join(lines)

    try:
        regulation = json.loads(content)
    except json.JSONDecodeError:
        match = re.search(r'\{[\s\S]*\}', content)
        if match:
            try:
                regulation = json.loads(match.group())
            except json.JSONDecodeError:
                return None
        else:
            return None
    return regulation


@app.route("/api/regulations/import/pdf", methods=["POST"])
def api_import_pdf():
    """
    上传PDF，创建异步解析任务。
    返回task_id，前端通过 /api/pdf/task/<task_id> 轮询进度。
    """
    if "file" not in request.files:
        return jsonify({"error": "请上传PDF文件"}), 400

    file = request.files["file"]
    if file.filename == "":
        return jsonify({"error": "未选择文件"}), 400

    if not file.filename.lower().endswith(".pdf"):
        return jsonify({"error": "仅支持PDF文件"}), 400

    file_bytes = file.read()
    mode = request.form.get("mode", "normal")
    use_ocr = mode == "ocr" or request.form.get("use_ocr", "false").lower() == "true"

    # 清理旧任务
    pdf_task_manager.cleanup_old_tasks()

    # 创建异步任务
    task_id = pdf_task_manager.create_task(file_bytes, file.filename, mode, use_ocr)

    return jsonify({
        "task_id": task_id,
        "message": "PDF解析任务已创建，请轮询进度",
        "filename": file.filename,
    })


@app.route("/api/pdf/task/<task_id>", methods=["GET"])
def api_pdf_task_status(task_id):
    """查询PDF解析任务进度"""
    task = pdf_task_manager.get_task(task_id)
    if not task:
        return jsonify({"error": "任务不存在或已过期"}), 404
    return jsonify(task)


@app.route("/api/pdf/task/<task_id>/cancel", methods=["POST"])
def api_pdf_task_cancel(task_id):
    """终止PDF解析任务"""
    success = pdf_task_manager.cancel_task(task_id)
    if success:
        return jsonify({"message": "任务已终止", "task_id": task_id})
    else:
        return jsonify({"error": "任务不存在或已完成，无法终止"}), 400


@app.route("/api/notice/review", methods=["POST"])
def api_review_notice():
    """AI二次复核：检查问题-条款对应关系及typo"""
    data = request.get_json()
    if not data:
        return jsonify({"error": "请提供通知书信息"}), 400

    project_info = data.get("project_info", {})
    problems = data.get("problems", [])

    if not problems:
        return jsonify({"error": "请至少提供一个问题"}), 400

    # 未配置LLM时直接通过
    if not llm.is_ready():
        return jsonify({
            "pass": True,
            "items": [],
            "summary": "未配置大模型API，跳过AI复核。",
            "llm_enabled": False,
        })

    # 如前端已传匹配结果，优先使用；否则重新匹配
    problem_results = []
    for prob in problems:
        desc = prob.get("description", "")
        matches = prob.get("matches") if "matches" in prob else matcher.match(desc, top_n=3)
        problem_results.append({
            "description": desc,
            "matches": matches,
        })

    review = llm.ai_review_notice(problem_results, project_info)
    if "error" in review:
        return jsonify({
            "pass": False,
            "error": review["error"],
            "items": [],
            "summary": review["error"],
            "llm_enabled": True,
        }), 500

    review["llm_enabled"] = True
    review["problems"] = problem_results
    return jsonify(review)


@app.route("/api/notice/generate", methods=["POST"])
def api_generate_notice():
    """生成整改通知书"""
    data = request.get_json()
    if not data:
        return jsonify({"error": "请提供通知书信息"}), 400

    project_info = data.get("project_info", {})
    problems = data.get("problems", [])
    skip_review = data.get("skip_review", False)

    # 如果前端已传匹配结果，直接使用；否则重新匹配
    problem_results = []
    for prob in problems:
        desc = prob.get("description", "")
        matches = prob.get("matches") if "matches" in prob else matcher.match(desc, top_n=3)
        corrected_desc = prob.get("corrected_description", desc)
        problem_results.append({
            "description": corrected_desc,
            "matches": matches,
        })

    # 二次复核（除非显式跳过）
    review_passed = skip_review or not llm.is_ready()
    review_result = None
    if not skip_review and llm.is_ready():
        review = llm.ai_review_notice(problem_results, project_info)
        if "error" not in review:
            review_passed = review.get("pass", False)
            review_result = review

    # 生成通知书文本
    notice_text = matcher.generate_notice_text(problem_results, project_info)

    # 保存记录
    conn = db._get_conn()
    conn.execute(
        """INSERT INTO notice_records
           (project_name, supervision_no, construction_unit,
            supervision_unit, construction_company, inspection_date, problems)
           VALUES (?, ?, ?, ?, ?, ?, ?)""",
        (
            project_info.get("project_name", ""),
            project_info.get("supervision_no", ""),
            project_info.get("construction_unit", ""),
            project_info.get("supervision_unit", ""),
            project_info.get("construction_company", ""),
            project_info.get("inspection_date", ""),
            json.dumps(problem_results, ensure_ascii=False),
        ),
    )
    conn.commit()
    conn.close()

    return jsonify({
        "notice_text": notice_text,
        "problems": problem_results,
        "review_passed": review_passed,
        "review_result": review_result,
    })


@app.route("/api/stats", methods=["GET"])
def api_stats():
    """获取数据库统计"""
    return jsonify(db.get_stats())


@app.route("/api/export/template", methods=["GET"])
def api_export_template():
    """导出规章制度导入模板"""
    template = {
        "regulations": [
            {
                "name": "示例规范名称",
                "code": "GB XXXXX-XXXX",
                "full_name": "示例规范全称",
                "category": "消防防火",
                "publish_date": "2022-01-01",
                "implement_date": "2022-07-01",
                "is_mandatory": True,
                "description": "规范简介",
                "clauses": [
                    {
                        "clause_number": "X.X.X",
                        "clause_content": "条款内容",
                        "keywords": ["关键词1", "关键词2"],
                        "category": "分类"
                    }
                ]
            }
        ]
    }
    return jsonify(template)


# ==============================================================================
# 通知书导出（Word/PDF）
# ==============================================================================
@app.route("/api/notice/export/word", methods=["POST"])
def api_export_notice_word():
    """导出整改通知书为 Word (.docx) 文件"""
    data = request.get_json()
    if not data or "notice_text" not in data:
        return jsonify({"error": "请提供通知书文本"}), 400

    notice_text = data["notice_text"]
    project_info = data.get("project_info", {})

    try:
        import io
        from docx import Document
        from docx.shared import Pt, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        # 设置默认字体
        style = doc.styles["Normal"]
        font = style.font
        font.name = "宋体"
        font.size = Pt(12)
        from docx.oxml.ns import qn
        style.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

        # 按行解析通知书文本
        lines = notice_text.split("\n")
        for line in lines:
            line = line.rstrip()
            if not line:
                # 空行
                doc.add_paragraph("")
                continue

            # 标题判断
            if "通知书" in line and len(line) < 30:
                p = doc.add_paragraph(line)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.size = Pt(18)
                    run.font.bold = True
                    run.font.name = "黑体"
                    run.element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
            elif line.startswith("监督注册号") or line.startswith("  监督注册号"):
                p = doc.add_paragraph(line.strip())
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                for run in p.runs:
                    run.font.size = Pt(11)
            elif "质量监督员" in line or "建设工程质量监督站" in line or "签收人" in line:
                p = doc.add_paragraph(line)
                for run in p.runs:
                    run.font.size = Pt(12)
            elif line.startswith("  ") or line.startswith("（"):
                p = doc.add_paragraph(line)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                # 编号问题条目
                if len(line) > 0 and line[0].isdigit() and "." in line[:3]:
                    p = doc.add_paragraph(line)
                    p.paragraph_format.first_line_indent = Cm(0.74)
                    for run in p.runs:
                        run.font.size = Pt(12)
                else:
                    p = doc.add_paragraph(line)
                    p.paragraph_format.first_line_indent = Cm(0.74)
                    for run in p.runs:
                        run.font.size = Pt(12)

        # 保存到内存
        output = io.BytesIO()
        doc.save(output)
        output.seek(0)

        filename = f"整改通知书_{project_info.get('project_name', '未命名')[:20]}.docx"

        return send_file(
            output,
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            as_attachment=True,
            download_name=filename,
        )
    except ImportError:
        return jsonify({"error": "Word导出依赖未安装，请运行 pip install python-docx"}), 500
    except Exception as e:
        return jsonify({"error": f"Word导出失败: {str(e)}"}), 500


@app.route("/api/notice/export/pdf", methods=["POST"])
def api_export_notice_pdf():
    """导出整改通知书为 PDF 文件"""
    data = request.get_json()
    if not data or "notice_text" not in data:
        return jsonify({"error": "请提供通知书文本"}), 400

    notice_text = data["notice_text"]
    project_info = data.get("project_info", {})

    try:
        import io
        from fpdf import FPDF

        pdf = FPDF(orientation="P", unit="mm", format="A4")
        pdf.add_page()
        pdf.set_auto_page_break(auto=True, margin=20)

        # 注册中文字体（使用系统字体）
        font_paths = [
            "C:/Windows/Fonts/simsun.ttc",
            "C:/Windows/Fonts/simhei.ttf",
            "C:/Windows/Fonts/msyh.ttc",
            "/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc",
            "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
        ]

        font_registered = False
        for fp in font_paths:
            try:
                pdf.add_font("CJK", "", fp, uni=True)
                pdf.add_font("CJK", "B", fp, uni=True)
                font_registered = True
                break
            except Exception:
                continue

        if not font_registered:
            # 回退：使用内置字体，中文可能显示为方框
            return jsonify({
                "error": "未找到中文字体文件，PDF导出需要系统中安装中文字体（如宋体simsum.ttc）"
            }), 500

        lines = notice_text.split("\n")

        for line in lines:
            line = line.rstrip()
            if not line:
                pdf.ln(3)
                continue

            # 标题
            if "通知书" in line and len(line) < 30:
                pdf.set_font("CJK", "B", 16)
                pdf.cell(0, 10, line, ln=True, align="C")
                pdf.ln(2)
            elif line.startswith("监督注册号") or line.strip().startswith("监督注册号"):
                pdf.set_font("CJK", "", 10)
                pdf.cell(0, 6, line.strip(), ln=True, align="R")
            elif "质量监督站" in line and "（盖章）" in line:
                pdf.set_font("CJK", "", 11)
                pdf.cell(0, 6, line.strip(), ln=True, align="C")
            elif "（盖章）" in line:
                pdf.set_font("CJK", "", 11)
                pdf.cell(0, 6, line.strip(), ln=True, align="C")
            elif line.strip().startswith("签收人"):
                pdf.set_font("CJK", "", 11)
                pdf.cell(0, 6, line.strip(), ln=True, align="L")
            else:
                # 正文
                pdf.set_font("CJK", "", 11)
                # 处理缩进
                stripped = line.strip()
                if stripped and stripped[0].isdigit() and "." in stripped[:3]:
                    # 编号条目
                    pdf.multi_cell(0, 6, "    " + stripped)
                else:
                    pdf.multi_cell(0, 6, "    " + stripped)
                pdf.ln(1)

        output = io.BytesIO()
        pdf.output(output)
        output.seek(0)

        filename = f"整改通知书_{project_info.get('project_name', '未命名')[:20]}.pdf"

        return send_file(
            output,
            mimetype="application/pdf",
            as_attachment=True,
            download_name=filename,
        )
    except ImportError:
        return jsonify({"error": "PDF导出依赖未安装，请运行 pip install fpdf2"}), 500
    except Exception as e:
        return jsonify({"error": f"PDF导出失败: {str(e)}"}), 500


# ==============================================================================
# LLM 相关路由
# ==============================================================================
@app.route("/api/llm/providers", methods=["GET"])
def api_llm_providers():
    """获取所有LLM提供商信息"""
    return jsonify({"providers": PROVIDER_PRESETS})


@app.route("/api/llm/config", methods=["GET"])
def api_llm_get_config():
    """获取当前LLM配置"""
    return jsonify(llm.get_config())


@app.route("/api/llm/config", methods=["POST"])
def api_llm_set_config():
    """设置LLM配置"""
    data = request.get_json()
    if not data:
        return jsonify({"error": "请提供配置数据"}), 400

    llm.update_config(
        provider=data.get("provider", ""),
        api_key=data.get("api_key", ""),
        model=data.get("model", ""),
        temperature=data.get("temperature", 0.3),
        max_tokens=data.get("max_tokens", 2048),
        enabled=data.get("enabled", False),
    )
    return jsonify({
        "message": "配置已保存",
        "config": llm.get_config(),
    })


@app.route("/api/llm/test", methods=["POST"])
def api_llm_test():
    """测试LLM连接"""
    data = request.get_json()

    # 如果请求中带了临时配置，先更新
    if data and (data.get("provider") or data.get("api_key") or data.get("model")):
        llm.update_config(
            provider=data.get("provider", llm.config["provider"]),
            api_key=data.get("api_key", llm.config["api_key"]),
            model=data.get("model", llm.config["model"]),
            enabled=True,
        )

    result = llm.test_connection()
    return jsonify(result)


@app.route("/api/llm/toggle", methods=["POST"])
def api_llm_toggle():
    """启用/禁用LLM"""
    data = request.get_json()
    enabled = data.get("enabled", False) if data else False
    llm.set_enabled(enabled)
    return jsonify({
        "enabled": llm.config["enabled"],
        "message": "LLM已启用" if enabled else "LLM已禁用",
    })


@app.route("/api/llm/match", methods=["POST"])
def api_llm_match():
    """AI增强匹配：使用LLM进行问题-规范匹配"""
    if not llm.is_ready():
        return jsonify({"error": "LLM未启用，请先在设置中配置并启用"}), 400

    data = request.get_json()
    if not data or "problem" not in data:
        return jsonify({"error": "请提供问题描述"}), 400

    problem = data["problem"].strip()
    location = data.get("location", "")

    full_problem = f"{problem}（位置：{location}）" if location else problem

    # 获取所有条款数据用于AI匹配
    clauses_data = db.get_all_clauses()

    result = llm.ai_match(full_problem, clauses_data)

    if "error" in result:
        return jsonify({"error": result["error"]}), 500

    return jsonify({
        "problem": problem,
        "location": location,
        "matches": result.get("matched", []),
        "analysis": result.get("analysis", ""),
        "usage": result.get("usage", {}),
        "total": len(result.get("matched", [])),
        "match_source": "ai",
    })


@app.route("/api/llm/analyze", methods=["POST"])
def api_llm_analyze():
    """AI问题分析：对现场问题进行深度分析"""
    if not llm.is_ready():
        return jsonify({"error": "LLM未启用，请先在设置中配置并启用"}), 400

    data = request.get_json()
    if not data or "problem" not in data:
        return jsonify({"error": "请提供问题描述"}), 400

    problem = data["problem"].strip()
    matched_clauses = data.get("matched_clauses", [])

    result = llm.ai_analyze(problem, matched_clauses)

    if "error" in result:
        return jsonify({"error": result["error"]}), 500

    return jsonify({
        "analysis": result.get("analysis", ""),
        "usage": result.get("usage", {}),
    })


@app.route("/api/llm/polish", methods=["POST"])
def api_llm_polish():
    """AI润色通知书"""
    if not llm.is_ready():
        return jsonify({"error": "LLM未启用，请先在设置中配置并启用"}), 400

    data = request.get_json()
    if not data or "notice_text" not in data:
        return jsonify({"error": "请提供通知书文本"}), 400

    notice_text = data["notice_text"]
    project_info = data.get("project_info", {})

    result = llm.ai_polish_notice(notice_text, project_info)

    if "error" in result:
        return jsonify({"error": result["error"]}), 500

    return jsonify({
        "polished": result.get("polished", ""),
        "usage": result.get("usage", {}),
    })


@app.route("/api/match/hybrid", methods=["POST"])
def api_match_hybrid():
    """混合匹配：先关键词匹配，再AI增强（如果LLM可用）"""
    data = request.get_json()
    if not data or "problem" not in data:
        return jsonify({"error": "请提供问题描述"}), 400

    problem = data["problem"].strip()
    location = data.get("location", "")
    top_n = data.get("top_n", 5)

    full_problem = f"{problem}（位置：{location}）" if location else problem

    # 第一步：关键词匹配
    keyword_matches = matcher.match(full_problem, top_n=top_n)

    result = {
        "problem": problem,
        "location": location,
        "keyword_matches": keyword_matches,
        "keyword_total": len(keyword_matches),
        "ai_enabled": llm.is_ready(),
    }

    # 第二步：如果LLM可用且关键词匹配度不高，使用AI增强
    use_ai = data.get("use_ai", True)
    if llm.is_ready() and use_ai:
        # 判断是否需要AI增强：关键词最佳匹配低于80%或用户强制要求
        need_ai = data.get("force_ai", False)
        if keyword_matches:
            best_score = keyword_matches[0].get("match_percentage", 0)
            if best_score < 80:
                need_ai = True
        else:
            need_ai = True

        if need_ai:
            clauses_data = db.get_all_clauses()
            ai_result = llm.ai_match(full_problem, clauses_data)
            if "matched" in ai_result:
                result["ai_matches"] = ai_result["matched"]
                result["ai_analysis"] = ai_result.get("analysis", "")
                result["ai_total"] = len(ai_result["matched"])
                result["ai_usage"] = ai_result.get("usage", {})

    return jsonify(result)


# ==============================================================================
# 启动
# ==============================================================================
if __name__ == "__main__":
    print("=" * 60)
    print("  娜娜的工程质量监督Agent")
    print("  请在浏览器中访问: http://127.0.0.1:5000")
    print("=" * 60)
    port = int(os.environ.get("PORT", 5000))
    host = os.environ.get("HOST", "0.0.0.0")
    app.run(host=host, port=port, debug=False)
